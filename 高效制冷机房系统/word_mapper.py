import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
import sys
import subprocess

class WordMapperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 格式自动映射工具")
        self.root.geometry("550x300")
        self.root.resizable(False, False)

        self.doc1_path = tk.StringVar()
        self.doc2_path = tk.StringVar()

        self._build_ui()

    def _build_ui(self):
        # 字体与边距配置
        pad_kwargs = {'padx': 10, 'pady': 10}

        # 模板文档 (Doc1)
        doc1_label = tk.Label(self.root, text="选择模板文档 (Doc1):", cursor="hand2")
        doc1_label.grid(row=0, column=0, sticky="e", **pad_kwargs)
        doc1_label.bind("<Button-1>", self._select_doc1)
        doc1_entry = tk.Entry(self.root, textvariable=self.doc1_path, width=40, state='readonly')
        doc1_entry.grid(row=0, column=1, **pad_kwargs)
        doc1_entry.bind("<Button-1>", self._select_doc1)
        tk.Button(self.root, text="浏览...", command=self._select_doc1).grid(row=0, column=2, **pad_kwargs)

        # 源文档 (Doc2)
        doc2_label = tk.Label(self.root, text="选择数据源文档 (Doc2):", cursor="hand2")
        doc2_label.grid(row=1, column=0, sticky="e", **pad_kwargs)
        doc2_label.bind("<Button-1>", self._select_doc2)
        doc2_entry = tk.Entry(self.root, textvariable=self.doc2_path, width=40, state='readonly')
        doc2_entry.grid(row=1, column=1, **pad_kwargs)
        doc2_entry.bind("<Button-1>", self._select_doc2)
        tk.Button(self.root, text="浏览...", command=self._select_doc2).grid(row=1, column=2, **pad_kwargs)

        # 执行映射按钮
        self.execute_btn = tk.Button(self.root, text="执行映射 (Execute)", command=self._start_execution, 
                                     width=20, bg="#4CAF50", fg="black")
        self.execute_btn.grid(row=2, column=0, columnspan=3, pady=15)

        # 状态回显控制台
        self.console = tk.Text(self.root, height=5, width=70, state='disabled', bg="#f0f0f0")
        self.console.grid(row=3, column=0, columnspan=3, padx=10, pady=5)
        self._log("系统就绪，等待操作...")

    def _select_doc1(self, event=None):
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            self.doc1_path.set(path)

    def _select_doc2(self, event=None):
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            self.doc2_path.set(path)

    def _log(self, message):
        """线程安全的日志输出"""
        def update_text():
            self.console.config(state='normal')
            self.console.insert(tk.END, message + "\n")
            self.console.see(tk.END)
            self.console.config(state='disabled')
        self.root.after(0, update_text)

    def _start_execution(self):
        doc1 = self.doc1_path.get()
        doc2 = self.doc2_path.get()

        # 安全与防御机制：输入校验
        if not doc1 or not doc2:
            messagebox.showwarning("校验失败", "请同时选择 Doc1 (模板) 和 Doc2 (数据源)！")
            return
        
        if not doc1.lower().endswith('.docx') or not doc2.lower().endswith('.docx'):
             messagebox.showwarning("校验失败", "仅支持 .docx 格式文件！")
             return

        self.execute_btn.config(state='disabled')
        
        # 清空控制台
        self.console.config(state='normal')
        self.console.delete(1.0, tk.END)
        self.console.config(state='disabled')
        
        self._log("执行中...")

        # 线程隔离：后台执行 IO 密集型操作
        thread = threading.Thread(target=self._process_mapping, args=(doc1, doc2), daemon=True)
        thread.start()

    def _process_mapping(self, doc1_path, doc2_path):
        try:
            try:
                from docx import Document
            except ModuleNotFoundError:
                self._log("异常：缺少依赖 python-docx。")
                self.root.after(
                    0,
                    lambda: messagebox.showerror(
                        "依赖缺失",
                        "未检测到 python-docx，请先安装后重试：\n\npip install python-docx",
                    ),
                )
                return

            self._log("步骤 1: 加载 Doc1 模板并清理现有段落...")
            template_doc = Document(doc1_path)
            
            # 清理模板原有的所有段落节点（仅保留样式）
            for p in template_doc.paragraphs:
                p_element = p._element
                parent = p_element.getparent()
                if parent is not None:
                    parent.remove(p_element)

            self._log("步骤 2: 挂载 Doc2 并提取数据...")
            source_doc = Document(doc2_path)
            
            self._log("步骤 3: 提取文本并映射样式至 Doc1 实例...")
            for p in source_doc.paragraphs:
                text = p.text
                style_name = p.style.name
                
                # 追加到 Doc1 实例中
                new_p = template_doc.add_paragraph(text)
                
                try:
                    # 获取模板中的对应样式
                    new_p.style = template_doc.styles[style_name]
                except KeyError:
                    # 异常拦截：如果找不到则降级为 'Normal'
                    try:
                        new_p.style = template_doc.styles['Normal']
                    except KeyError:
                        pass # 极端情况：连 Normal 样式都没有，则保持默认

            # 步骤 4: 编译输出
            self._log("步骤 4: 正在编译并输出文件...")
            output_dir = os.path.dirname(doc2_path)
            output_filename = "Output_Formatted.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            # IO 防御：由于保存文件可能覆盖被占用的文件，由 try-except 的 PermissionError 拦截
            template_doc.save(output_path)
            
            self._log(f"成功！文件已生成: {output_filename}")
            self.root.after(0, lambda: messagebox.showinfo("成功", f"格式映射完毕！\n保存位置:\n{output_path}"))

        except PermissionError:
            self._log("异常：输出文件被其他程序占用！")
            self.root.after(0, lambda: messagebox.showerror("IO 错误", "文件权限被拒绝 (PermissionError)！\n请检查「Output_Formatted.docx」是否已在 Word 中打开，关闭后再试。"))
        except Exception as e:
            self._log(f"异常：{str(e)}")
            self.root.after(0, lambda: messagebox.showerror("执行错误", f"发生未预期的错误:\n{str(e)}"))
        finally:
            self.root.after(0, lambda: self.execute_btn.config(state='normal'))

def _bootstrap_macos_tk():
    """
    兼容部分 macOS 场景下 "python script.py" 非交互启动导致的 Tk 崩溃：
    自动重启为 stdin 启动方式，避免 NSMenuItem 初始化异常。
    """
    if sys.platform != "darwin":
        return
    if os.environ.get("WORD_MAPPER_TK_BOOTSTRAPPED") == "1":
        return
    script_path = globals().get("__file__")
    if not script_path or not os.path.isfile(script_path):
        return
    if sys.stdout.isatty():
        return

    with open(script_path, "r", encoding="utf-8") as f:
        script_source = f.read()

    env = os.environ.copy()
    env["WORD_MAPPER_TK_BOOTSTRAPPED"] = "1"
    result = subprocess.run([sys.executable, "-"], input=script_source, text=True, env=env)
    raise SystemExit(result.returncode)

if __name__ == "__main__":
    _bootstrap_macos_tk()
    root = tk.Tk()
    app = WordMapperApp(root)
    root.mainloop()
